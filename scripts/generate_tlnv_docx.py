from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def parse_markdown(md_path: Path, docx_path: Path):
    content = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    # Title
    doc.add_heading("Team Thesis Documentation (TLNV) - IScream Project", level=0)

    # TOC
    toc_title = doc.add_paragraph()
    toc_title.add_run("Table of Contents").bold = True
    toc_para = doc.add_paragraph()
    add_toc(toc_para)

    in_code = False
    code_buffer = []

    for line in content:
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                if code_buffer:
                    p = doc.add_paragraph("\n".join(code_buffer))
                    p.style = doc.styles["No Spacing"]
                    for run in p.runs:
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        # Keep numbered and roman items as plain paragraph to preserve original text.
        doc.add_paragraph(line)

    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "TLNV_IScream.md"
    out = root / "TLNV_IScream.docx"
    parse_markdown(md, out)
    print(str(out))
